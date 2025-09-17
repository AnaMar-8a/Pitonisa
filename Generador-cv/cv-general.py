from docx import Document

# ===============================================================
# PLANTILLA DE CV PROFESIONAL EN BLANCO
# ---------------------------------------------------------------
# Este script genera un documento Word (.docx) con la estructura
# básica de un CV profesional. Incluye secciones comunes como:
# Perfil Profesional, Experiencia, Educación, Certificaciones y
# Habilidades. Cada sección está vacía para que sea personalizada.
# ===============================================================

# Crear documento
doc = Document()

# Encabezado
doc.add_heading("NOMBRE COMPLETO", level=0)
doc.add_paragraph("Cargo objetivo | Industria / Especialidad")
doc.add_paragraph("Ciudad, País | email@example.com | +57 000 000 0000 | linkedin.com/in/usuario")

# Perfil Profesional
doc.add_heading("Perfil Profesional", level=1)
doc.add_paragraph(
    "Breve resumen (4-6 líneas) de tu experiencia, conocimientos clave, "
    "logros relevantes y lo que buscas aportar en un nuevo rol."
)

# Experiencia Profesional
doc.add_heading("Experiencia Profesional", level=1)

doc.add_heading("Cargo | Empresa", level=2)
doc.add_paragraph("Ciudad, País | Mes Año – Mes Año")
doc.add_paragraph("• Responsabilidad o tarea principal")
doc.add_paragraph("• Responsabilidad o tarea secundaria")
doc.add_paragraph("✔ Logro destacado (ejemplo: reducción de costos en 20%)")

doc.add_paragraph("")  # Espacio

# Educación
doc.add_heading("Educación", level=1)
doc.add_paragraph("Título | Universidad | Año")

# Certificaciones
doc.add_heading("Certificaciones", level=1)
doc.add_paragraph("Certificación | Año | Institución")

# Habilidades
doc.add_heading("Habilidades", level=1)
doc.add_paragraph(
    "• Habilidad 1\n"
    "• Habilidad 2\n"
    "• Habilidad 3"
)

# Guardar documento
doc.save("CV_Plantilla_Blanco.docx")
print("Plantilla de CV generada: CV_Plantilla_Blanco.docx")
