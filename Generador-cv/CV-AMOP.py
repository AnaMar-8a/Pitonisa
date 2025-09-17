from docx import Document

# ===============================================================
# GENERADOR DE CV PROFESIONAL (DOCX)
# Autor: Ana María Ochoa Patiño (con asistencia de IA)
# Descripción:
#   Este script construye un CV profesional completo y actualizado
#   utilizando la librería python-docx. Está organizado en secciones
#   claras (Perfil, Experiencia, Educación, Certificaciones, Habilidades),
#   e incluye buenas prácticas de legibilidad y reutilización.
# ===============================================================

# ---------------------------------------------------------------
# Función para agregar una sección de experiencia laboral
# ---------------------------------------------------------------
def add_experience(doc, title, meta, bullets, achievements=None):
    """
    Agrega una experiencia profesional al documento.
    :param doc: objeto Document
    :param title: cargo y empresa
    :param meta: ubicación y fechas
    :param bullets: lista de responsabilidades principales
    :param achievements: lista de logros destacados (opcional)
    """
    doc.add_heading(title, level=2)
    doc.add_paragraph(meta)
    for b in bullets:
        doc.add_paragraph(f"• {b}")
    if achievements:
        doc.add_paragraph("\nLogros destacados:")
        for a in achievements:
            doc.add_paragraph(f"✔ {a}")
    doc.add_paragraph("")


# ---------------------------------------------------------------
# Crear documento
# ---------------------------------------------------------------
doc = Document()

# Encabezado
doc.add_heading("Ana María Ochoa Patiño", level=0)
doc.add_paragraph("Senior Project Manager | Telecomunicaciones | Transformación Digital | Software & Datos")
doc.add_paragraph("Medellín, Colombia | 8a.anamaria@gmail.com | +57 311 326 26 20 | linkedin.com/in/8aanamaria")

# Perfil Profesional
doc.add_heading("Perfil Profesional", level=1)
doc.add_paragraph(
    "Senior Project Manager con más de 8 años de experiencia gestionando proyectos en telecomunicaciones, software y transformación digital "
    "en entornos multicountry (Colombia, México, Chile y Perú). Experta en la planificación, ejecución y cierre de proyectos de alta complejidad, "
    "desde despliegues de infraestructura de telecomunicaciones hasta implementaciones de software corporativo en la nube.\n\n"
    "He liderado equipos globales y remotos en proyectos de SaaS, ERP (SAP S4/HANA), gobierno de datos y tendido de fibra óptica de más de 500 km, "
    "trabajando con partners como Claro, TIGO y Ruta N. Me especializo en gestión integral de proyectos (alcance, cronograma, presupuesto, riesgos, dependencias), "
    "normatividad técnica (RETIE, RITEL, HSEQ), cultura y gobierno de datos (Power BI, ISO 8000) y metodologías ágiles e híbridas (Scrum, Kanban, Waterfall)."
)

# ---------------------------------------------------------------
# Sección: Experiencia Profesional
# ---------------------------------------------------------------
doc.add_heading("Experiencia Profesional", level=1)

# Cada experiencia se agrega con la función add_experience
add_experience(
    doc,
    "Consultora Senior de Transformación Digital | Grupo GTA",
    "Colombia, Chile, Perú | Ene 2024 – Presente",
    [
        "Lideré la implementación multinacional de SAP S4/HANA en cloud (RISE with SAP), gestionando un presupuesto de $1B COP y coordinando equipos interdisciplinarios en tres países.",
        "Actúo como enlace entre stakeholders de negocio y tecnología, elaborando y presentando reportes ejecutivos de avance, riesgos y dependencias a C-Level para la toma de decisiones estratégicas.",
        "Diseñé el Plan Estratégico de Tecnología (PETI), alineando 15 iniciativas TI con los OKRs corporativos.",
        "Implementé un modelo de gobernanza ágil para equipos remotos, incrementando en 30% la adopción de nuevas prácticas y eficiencia operativa.",
    ],
    [
        "Ejecución multinacional sin desviaciones críticas en presupuesto ni cronograma.",
        "Alineación estratégica de TI con los objetivos corporativos, habilitando proyectos de innovación digital.",
    ],
)

add_experience(
    doc,
    "Directora de Proyectos | Syspotec",
    "Medellín, Colombia | Ene 2022 – Ene 2024",
    [
        "Gestioné simultáneamente un portafolio de 25 proyectos de desarrollo de software para 6 clientes clave del sector utilities, con un presupuesto anual de $3.000M COP.",
        "Lideré y mentoricé 6 squads ágiles (~30 personas) en un modelo de fábrica de software, optimizando la productividad y calidad de entregas.",
        "Validé requerimientos con stakeholders estratégicos y presenté informes ejecutivos de riesgos y dependencias a gerencia y C-Level.",
        "Establecí KPIs de calidad que redujeron en 40% los bugs reportados en producción en 12 meses.",
    ],
    [
        "Consolidación de un modelo de fábrica de software eficiente y escalable.",
        "Incremento de la satisfacción de clientes estratégicos gracias a entregas de mayor calidad.",
    ],
)

add_experience(
    doc,
    "Líder de Transformación Digital | AIR-E (Sector Energético)",
    "Barranquilla, Colombia | Oct 2020 – Dic 2022",
    [
        "Diseñé y ejecuté la estrategia de cultura ágil, capacitando a más de 200 colaboradores en Scrum y Kanban a través de 12 workshops.",
        "Creé el equipo de analítica corporativa y el programa 'Data Champions', estableciendo un gobierno de datos basado en la norma ISO 8000.",
        "Centralicé indicadores de gestión en Power BI, habilitando una toma de decisiones más rápida y confiable.",
        "Implementé un modelo de feedback 360° para alinear proyectos estratégicos con las necesidades internas.",
    ],
    [
        "Incremento del NPS interno de 50 a 85/100.",
        "Creación de una cultura de datos sostenible que aumentó la confiabilidad y trazabilidad de la información.",
    ],
)

add_experience(
    doc,
    "Product Owner | Digita Studio (Startup Mexicana)",
    "Ciudad de México, México | Feb 2022 – Dic 2022",
    [
        "Gestioné el backlog y roadmap de un producto SaaS regional para gestión operativa en 4 países LATAM.",
        "Definí user stories, épicas y criterios de aceptación; coordiné desarrollo, QA y UX/UI en Scrum.",
        "Apliqué MoSCoW y RICE para priorización, logrando 90% de cumplimiento en entregas planificadas.",
    ],
    [
        "Lanzamiento exitoso del producto SaaS en 4 países, adaptado a distintos contextos regulatorios.",
        "Reducción del retrabajo en 25% mediante validación temprana con usuarios.",
    ],
)

add_experience(
    doc,
    "Gestora de Proyectos TI | Indra (Cliente TIGO)",
    "Medellín, Colombia | Ago 2019 – May 2020",
    [
        "Indra, consultora multinacional de tecnología. Gestioné proyectos estratégicos para TIGO, apoyando la mejora de procesos y la automatización con RPA (UiPath).",
        "Gestioné backlog de solicitudes y mejoras en Jira; definí y documenté user stories, casos de uso y criterios de aceptación.",
        "Coordiné UX/UI, QA y DevOps, asegurando integración continua y despliegues estables.",
        "Lideré transferencia de conocimiento (KT), reduciendo en 30% la dependencia del cliente hacia equipos externos.",
    ],
    [
        "Implementación de automatizaciones que optimizaron procesos de soporte.",
        "Mejora de la estabilidad de entregas en entornos productivos.",
    ],
)

add_experience(
    doc,
    "Instructora – Área de Telecomunicaciones | SENA",
    "Medellín, Colombia | Nov 2016 – Ago 2019",
    [
        "Planifiqué y ejecuté procesos formativos en telecomunicaciones, garantizando coherencia entre programa y perfiles de aprendices.",
        "Coordiné actividades prácticas en laboratorios especializados en fibra óptica, promoviendo competencias aplicadas para despliegue de redes.",
        "Articulé formación con mesas sectoriales de telecomunicaciones para alinear programas a demanda del mercado.",
    ],
    [
        "Formación de más de 400 aprendices y fortalecimiento de la empleabilidad técnica en el sector.",
    ],
)

add_experience(
    doc,
    "Coordinadora de Formación y Proyectos de Telecomunicaciones | CYFO",
    "Medellín, Colombia | Ene 2015 – Oct 2016",
    [
        "Diseñé y ejecuté el programa nacional de formación técnica en telecomunicaciones e innovación.",
        "Gestioné convenios con instituciones educativas y técnicas y coordiné proyectos de tendido de fibra óptica de más de 500 km en alianza con Claro.",
        "Supervisé planificación de recursos, logística y proveedores para proyectos de infraestructura.",
    ],
    [
        "Implementación de programa formativo que mejoró productividad y retención de personal técnico.",
        "Ejecución exitosa de despliegues de fibra óptica a gran escala.",
    ],
)

add_experience(
    doc,
    "Especialista en Proyectos | Universidad de Antioquía",
    "Medellín, Colombia | 2018",
    [
        "Apoyé la construcción del plan de negocios para el desarrollo de un protocolo de terapia celular avanzada en alianza entre Ruta N y el Grupo de Ingeniería de Tejidos y Terapias Celulares.",
        "Coordiné reuniones entre investigadores, líderes de innovación y stakeholders para viabilizar el proyecto.",
    ],
    [
        "Estructuración de un plan de negocio que facilitó la proyección hacia validación clínica y transferencia tecnológica.",
    ],
)

add_experience(
    doc,
    "Docente de Cátedra | Politécnico Jaime Isaza Cadavid",
    "Medellín, Colombia | Ene 2023 – Presente",
    [
        "Imparto asignaturas en telecomunicaciones, comunicaciones inalámbricas, fibra óptica e innovación.",
        "Diseño de planes de clase, prácticas en laboratorio y metodologías activas de aprendizaje.",
        "Incorporación de conceptos de innovación tecnológica en el currículo académico.",
    ],
)

# ---------------------------------------------------------------
# Sección: Educación
# ---------------------------------------------------------------
doc.add_heading("Educación", level=1)
doc.add_paragraph("Especialización en Evaluación de Proyectos | Universidad de Antioquía | 2019")
doc.add_paragraph("Ingeniera de Telecomunicaciones | Universidad Santo Tomás | 2014")

# ---------------------------------------------------------------
# Sección: Certificaciones
# ---------------------------------------------------------------
doc.add_heading("Certificaciones", level=1)
doc.add_paragraph(
    "Agile Coach Certified Expert (ACCE)\n"
    "Scrum Master Certified Expert (SMCE)\n"
    "Scrum Product Owner Certified Expert (SPOCE)\n"
    "Kanban Certified Associate (KCA)\n"
    "OKR Certified Associate (OCA)\n"
    "Remote Worker Professional Certificate (RWPC)"
)

# ---------------------------------------------------------------
# Sección: Habilidades Técnicas
# ---------------------------------------------------------------
doc.add_heading("Habilidades Técnicas", level=1)
doc.add_paragraph(
    "Gestión de proyectos: SDLC, alcan
